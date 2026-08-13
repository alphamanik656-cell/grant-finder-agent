import re
from datetime import datetime
from pathlib import Path

_BASE       = Path(__file__).parent.parent / "output"
REPORTS_DIR = _BASE / "reports"
DRAFTS_DIR  = _BASE / "drafts"


# ── Markdown block parser ─────────────────────────────────────────────────────

def _parse_blocks(md: str) -> list[dict]:
    blocks = []
    for line in md.splitlines():
        if line.startswith("#### "):
            blocks.append({"type": "h4", "text": line[5:]})
        elif line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:]})
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:]})
        elif line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:]})
        elif line.startswith("> "):
            blocks.append({"type": "quote", "text": line[2:]})
        elif re.match(r"^[-*] ", line):
            blocks.append({"type": "bullet", "text": line[2:]})
        elif re.match(r"^\d+\. ", line):
            blocks.append({"type": "num", "text": re.sub(r"^\d+\.\s+", "", line)})
        elif line.strip() == "---":
            blocks.append({"type": "hr"})
        elif line.strip() == "":
            blocks.append({"type": "blank"})
        else:
            blocks.append({"type": "para", "text": line})
    return blocks


# ── Inline formatting ─────────────────────────────────────────────────────────

_INLINE_PAT = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _split_inline(text: str) -> list[dict]:
    result = []
    for part in _INLINE_PAT.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            result.append({"t": part[2:-2], "bold": True})
        elif part.startswith("*") and part.endswith("*"):
            result.append({"t": part[1:-1], "italic": True})
        elif part.startswith("`") and part.endswith("`"):
            result.append({"t": part[1:-1], "code": True})
        elif m := re.match(r"\[([^\]]+)\]\(([^)]+)\)", part):
            result.append({"t": m.group(1), "url": m.group(2)})
        else:
            result.append({"t": part})
    return result


def _rl_inline(text: str) -> str:
    """Convert Markdown inline to ReportLab XML markup."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*([^*]+)\*",     r"<i>\1</i>", text)
    text = re.sub(r"`([^`]+)`",       r'<font name="Courier">\1</font>', text)
    text = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        r'<a href="\2"><u><font color="blue">\1</font></u></a>',
        text,
    )
    text = re.sub(
        r'(?<!["\'])https?://[^\s<>&"\']+',
        lambda m: f'<a href="{m.group()}">{m.group()}</a>',
        text,
    )
    return text


# ── Word (.docx) export ───────────────────────────────────────────────────────

def _add_hyperlink_docx(para, text: str, url: str):
    """Add a proper clickable hyperlink to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    # Use the built-in Hyperlink character style (blue + underline)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    # Explicit colour + underline as fallback if the style isn't in the template
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)

    hl.append(r)
    para._p.append(hl)


def _add_inline_docx(para, text: str):
    for seg in _split_inline(text):
        if seg.get("url"):
            _add_hyperlink_docx(para, seg["t"], seg["url"])
        else:
            run = para.add_run(seg["t"])
            if seg.get("bold"):
                run.bold = True
            if seg.get("italic"):
                run.italic = True
            if seg.get("code"):
                run.font.name = "Courier New"


def _md_to_docx(doc, md: str):
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_hr():
        para = doc.add_paragraph()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        para._p.get_or_add_pPr().append(pBdr)

    for block in _parse_blocks(md):
        btype = block["type"]
        text  = block.get("text", "")

        if btype in ("h1", "h2", "h3", "h4"):
            level = int(btype[1])
            p = doc.add_paragraph(style=f"Heading {level}")
            _add_inline_docx(p, text)
        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_docx(p, text)
        elif btype == "num":
            p = doc.add_paragraph(style="List Number")
            _add_inline_docx(p, text)
        elif btype == "quote":
            try:
                p = doc.add_paragraph(style="Quote")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
            _add_inline_docx(p, text)
        elif btype == "hr":
            add_hr()
        elif btype == "blank":
            pass
        else:
            if text.strip():
                p = doc.add_paragraph()
                _add_inline_docx(p, text)


def _make_docx(content: str) -> "Document":
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    for section in doc.sections:
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    _md_to_docx(doc, content)
    return doc


def save_report_docx(content: str, filename: str) -> str:
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stem = re.sub(r"\.(md|docx)$", "", filename)
    path = REPORTS_DIR / f"{stem}.docx"
    _make_docx(content).save(str(path))
    return str(path)


def save_draft_docx(content: str, grant_name: str) -> str:
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^a-z0-9]+", "-", grant_name.lower()).strip("-")
    date = datetime.now().strftime("%Y-%m-%d")
    path = DRAFTS_DIR / f"draft-{safe}-{date}.docx"
    _make_docx(content).save(str(path))
    return str(path)


# ── PDF export ────────────────────────────────────────────────────────────────

def _md_to_pdf_story(md: str) -> list:
    from reportlab.lib.pagesizes import letter  # noqa: F401
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, HRFlowable

    base = getSampleStyleSheet()

    H1   = ParagraphStyle("rH1", parent=base["Heading1"],  fontSize=18, spaceAfter=10, spaceBefore=4,  textColor=colors.HexColor("#1e2233"))
    H2   = ParagraphStyle("rH2", parent=base["Heading2"],  fontSize=14, spaceAfter=8,  spaceBefore=14, textColor=colors.HexColor("#2c3e50"))
    H3   = ParagraphStyle("rH3", parent=base["Heading3"],  fontSize=12, spaceAfter=6,  spaceBefore=10, textColor=colors.HexColor("#34495e"))
    H4   = ParagraphStyle("rH4", parent=base["Normal"],    fontSize=11, spaceAfter=4,  spaceBefore=8,  textColor=colors.HexColor("#444444"), fontName="Helvetica-Bold")
    BODY = ParagraphStyle("rBD", parent=base["Normal"],    fontSize=10, leading=15,    spaceAfter=6)
    BUL  = ParagraphStyle("rBL", parent=base["Normal"],    fontSize=10, leading=15,    spaceAfter=3,   leftIndent=18)
    QT   = ParagraphStyle("rQT", parent=base["Normal"],    fontSize=10, leading=15,    leftIndent=30,  textColor=colors.grey)

    story      = []
    buf_bullets = []

    def flush():
        for bt in buf_bullets:
            story.append(Paragraph(f"• &nbsp;{bt}", BUL))
        buf_bullets.clear()

    for block in _parse_blocks(md):
        btype = block["type"]
        text  = _rl_inline(block.get("text", ""))

        if btype != "bullet":
            flush()

        if btype == "h1":
            story.append(Paragraph(text, H1))
        elif btype == "h2":
            story.append(Spacer(1, 4))
            story.append(Paragraph(text, H2))
        elif btype == "h3":
            story.append(Paragraph(text, H3))
        elif btype == "h4":
            story.append(Paragraph(f"<b>{text}</b>", H4))
        elif btype == "bullet":
            buf_bullets.append(text)
        elif btype == "num":
            story.append(Paragraph(f"• &nbsp;{text}", BUL))
        elif btype == "quote":
            story.append(Paragraph(text, QT))
        elif btype == "hr":
            story.append(Spacer(1, 6))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
            story.append(Spacer(1, 6))
        elif btype == "blank":
            story.append(Spacer(1, 4))
        else:
            if text.strip():
                story.append(Paragraph(text, BODY))

    flush()
    return story


def _make_pdf(content: str, out_path: Path):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.units import inch
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        leftMargin=1.2 * inch, rightMargin=1.2 * inch,
        topMargin=1.0 * inch,  bottomMargin=1.0 * inch,
    )
    doc.build(_md_to_pdf_story(content))


def save_report_pdf(content: str, filename: str) -> str:
    try:
        import reportlab  # noqa: F401
    except ImportError:
        raise RuntimeError("reportlab not installed — run: pip install reportlab")
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stem = re.sub(r"\.(md|pdf)$", "", filename)
    path = REPORTS_DIR / f"{stem}.pdf"
    _make_pdf(content, path)
    return str(path)


def save_draft_pdf(content: str, grant_name: str) -> str:
    try:
        import reportlab  # noqa: F401
    except ImportError:
        raise RuntimeError("reportlab not installed — run: pip install reportlab")
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^a-z0-9]+", "-", grant_name.lower()).strip("-")
    date = datetime.now().strftime("%Y-%m-%d")
    path = DRAFTS_DIR / f"draft-{safe}-{date}.pdf"
    _make_pdf(content, path)
    return str(path)


# ── In-memory bytes (for Streamlit Cloud download buttons) ───────────────────

def report_as_docx_bytes(content: str) -> bytes:
    """Return the report as a .docx file in memory (no disk write)."""
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    import io
    buf = io.BytesIO()
    _make_docx(content).save(buf)
    return buf.getvalue()


def report_as_pdf_bytes(content: str) -> bytes:
    """Return the report as a PDF file in memory (no disk write)."""
    try:
        import reportlab  # noqa: F401
    except ImportError:
        raise RuntimeError("reportlab not installed — run: pip install reportlab")
    import io
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.units import inch
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=1.2 * inch, rightMargin=1.2 * inch,
        topMargin=1.0 * inch,  bottomMargin=1.0 * inch,
    )
    doc.build(_md_to_pdf_story(content))
    return buf.getvalue()


# ── Markdown saves (original) ─────────────────────────────────────────────────

def save_report(content: str, filename: str) -> str:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    if not filename.endswith(".md"):
        filename = f"{filename}.md"
    path = REPORTS_DIR / filename
    path.write_text(content, encoding="utf-8")
    return str(path)


def save_draft(content: str, grant_name: str) -> str:
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^a-z0-9]+", "-", grant_name.lower()).strip("-")
    date = datetime.now().strftime("%Y-%m-%d")
    path = DRAFTS_DIR / f"draft-{safe}-{date}.md"
    path.write_text(content, encoding="utf-8")
    return str(path)
